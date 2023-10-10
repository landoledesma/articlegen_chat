from langchain.prompts import PromptTemplate
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
import os
import openai
from dotenv import load_dotenv
import streamlit as st 
import os
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import requests
from fetch_images import fetch_photo

def load_llm(template):
    llm  = ChatOpenAI(
        max_tokens=2000,
        temperature=0.5
        )
    
    llm_chain = LLMChain(
        llm = llm,
        prompt=PromptTemplate.from_template(template),
    )
    return llm_chain

query = 'kitties'
src_original_url = fetch_photo(query)
if src_original_url:
    print(f"Original URL for query '{query}': {src_original_url}")
else:
    print("No se pudo traer imagens lo siento")

def create_docx(user_input,paragraph,image_input):
    doc = Document()
    doc.add_heading(user_input)
    doc.add_paragraph(paragraph)

    doc.add_heading('Image',level=1)
    image_stream = io.BytesIO()
    image_input.save(image_stream,format='PNG')
    image_input.seek(0)
    doc.add_picture(image_stream,width=Inches(4))

    return doc

st.set_page_config(layout="wide")

def main():
    st.title("Aplicación Generadora de Artículoscon AI")

    entrada_usuario = st.text_input("¡Por favor, introduce la idea/tema para el artículo que deseas generar!")

    entrada_imagen = st.text_input("¡Por favor, introduce el tema de la imagen que deseas obtener!")

    if len(entrada_usuario) > 0 and len(entrada_imagen) > 0:

        col1, col2, col3 = st.columns([1,2,1])

        with col1:
            st.subheader("Contenido Generado con inteligencia artifical")
            st.write("El tema del artículo es: " + entrada_usuario)
            st.write("La imagen del artículo es: " + entrada_imagen)
            plantilla_prompt = """Eres un experto en marketing digital y SEO y tu tarea es escribir un artículo
              sobre el tema proporcionado: {entrada_usuario}. El artículo no debe superar las 800 palabras.
              El artículo no debe ser muy largo.
            """
            llamada_llm = cargar_llm(max_tokens=800, plantilla_prompt=plantilla_prompt)
            print(llamada_llm)
            resultado = llamada_llm(entrada_usuario)
            if len(resultado) > 0:
                st.info("¡Tu artículo ha sido generado con éxito!")
                st.write(resultado)
            else:
                st.error("¡No se pudo generar tu artículo!")

        with col2:
            st.subheader("Imagen Obtenida")
            url_imagen = obtener_url_original(entrada_imagen)
            st.image(url_imagen)


        with col3:
            st.subheader("Articulo final")
            image_input = "temp_image.jpg"
            doc = create_word_docx(user_input, result['text'], Image.open(image_input))

            # Save the Word document to a BytesIO buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Prepare the download link
            st.download_button(
                label='Download Word Document',
                data=doc_buffer,
                file_name='document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )


if __name__ == "__main__":
    main()
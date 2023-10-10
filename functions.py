from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import os
from dotenv import load_dotenv
import requests

load_dotenv("token.env")
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

def load_llm(template):
    llm  = ChatOpenAI(
        max_tokens=2000,
        temperature=0.5
    )
    
    llm_chain = LLMChain(
        llm = llm,
        prompt=PromptTemplate.from_template(template),
    )
    return llm_chain

def create_docx(user_input, paragraph, image_input):
    doc = Document()
    doc.add_heading(user_input)
    doc.add_paragraph(paragraph)
    doc.add_heading('Image', level=1)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(4))
    return doc


